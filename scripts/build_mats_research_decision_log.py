"""Build the source-backed research rationale and pivot history as a DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "mats_application" / "multimodal_jspace_research_decision_log_2026-08-19.docx"
FIG = ROOT / "reports" / "mats_application" / "figures"

BLACK = "202124"
GRAY = "5F6368"
LIGHT = "F8F9FA"
GRID = "DADCE0"
BLUE = "1A73E8"
PALE_BLUE = "E8F0FE"
GREEN = "188038"
PALE_GREEN = "E6F4EA"
AMBER = "B06000"
PALE_AMBER = "FEF7E0"
RED = "C5221F"
PALE_RED = "FCE8E6"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_table_together(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True


def configure(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after, color in (
        ("Heading 1", 20, 20, 6, BLACK),
        ("Heading 2", 16, 18, 6, BLACK),
        ("Heading 3", 14, 16, 4, GRAY),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(26)
    r.font.bold = False
    r.font.color.rgb = RGBColor.from_string(BLACK)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_callout(doc: Document, label: str, text: str, kind: str = "blue") -> None:
    fill, color = {
        "blue": (PALE_BLUE, BLUE),
        "green": (PALE_GREEN, GREEN),
        "amber": (PALE_AMBER, AMBER),
        "red": (PALE_RED, RED),
        "gray": (LIGHT, GRAY),
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(color)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = widths is None
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "E8EAED")
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(BLACK)
        margins(cell)
        if widths:
            cell.width = Inches(widths[i])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 == 0 else LIGHT
        for i, value in enumerate(values):
            shade(cells[i], fill)
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.name = "Arial"
            r.font.size = Pt(8.6)
            r.font.color.rgb = RGBColor.from_string(BLACK)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cells[i])
            if widths:
                cells[i].width = Inches(widths[i])
    keep_table_together(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(doc: Document, filename: str, caption: str, width=6.45) -> None:
    path = FIG / filename
    if not path.is_file():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(caption)
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def add_source(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Evidence: ")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_explicit_h1(doc: Document, text: str) -> None:
    """Add a zero-indented H1 after tables that may leak paragraph geometry."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.bold = False
    r.font.color.rgb = RGBColor.from_string(BLACK)


PIVOTS = [
    {
        "title": "1. Port the method to Gemma 4 before asking a multimodal question",
        "belief": "The project first needed an architecture-correct J-lens implementation. A multimodal result would be uninterpretable if the layer path, final normalization, unembedding, softcap, or tied weights were wrong.",
        "trigger": "The starting repository targeted the public Global Workspace method, not Gemma 4 E4B’s exact module layout.",
        "decision": "Build and test a Gemma 4 adapter, freeze model weights, record 42 layers, d_model=2,560, vocab=262,144, tied unembedding, final-logit softcap, layer scalars, and the block-output hook convention.",
        "outcome": "The adapter and audit path became the common foundation for every later run. The layer-scalar warning was recorded as architectural metadata rather than treated as fatal.",
        "claim": "Engineering readiness only; no scientific claim.",
        "source": "commits b5953c7 and 348405b; audio_audit.json; model revision fa62d88…",
    },
    {
        "title": "2. Make long Colab work resumable and provenance-bound",
        "belief": "Runs lasting hours needed to survive disconnects without silently mixing configurations.",
        "trigger": "The fitting and causal budgets were too large for an all-or-nothing notebook session.",
        "decision": "Introduce atomic unit files, rolling accumulators, run fingerprints, Drive-backed directories, immutable parent checks, and refusal on changed configurations.",
        "outcome": "Later 250- and 1,000-prompt fits resumed from checkpoints; large causal studies reused checksum-valid units.",
        "claim": "Reliability improvement. Resumability does not make a result valid, but it prevents torn or mixed runs.",
        "source": "commits ca33c54 and d62814f; fingerprint.json and resume blocks across archived runs",
    },
    {
        "title": "3. Establish a text-only calibration baseline",
        "belief": "A population-average Jacobian estimated on general text should provide a reusable verbalizable coordinate system.",
        "trigger": "Before adding images or speech, the team needed to know where the lens was faithful enough to read the model.",
        "decision": "Fit the running-average Jacobian on 100 WikiText-103 prompts and validate multiple layers with native-readout, noise, shuffled, wrong-layer, rank, and stability controls.",
        "outcome": "Untouched confirmation passed at L35, L38, and L40; L32 did not pass at scale 100. Increasing fit scale later did not materially improve the eligible layer set.",
        "claim": "Validated text-calibrated lenses existed at late layers. This did not yet establish multimodal transfer.",
        "source": "rgcalib…/artifacts/calibration_report.json; commit 0548c05; report checksum sha256:db71…",
    },
    {
        "title": "4. Probe whether sparse J-space structure was nontrivial",
        "belief": "If the lens captured useful coordinates, sparse decompositions should show structured support across nearby positions and layers.",
        "trigger": "A calibrated matrix alone did not show whether individual activations had interpretable sparse structure.",
        "decision": "Run gradient-pursuit decompositions at five layers and three sparsity levels, with integrity checks and transition comparisons.",
        "outcome": "1,140 decompositions were analyzed. Exact signature repeats were zero; explained fractions were small, and L21 was an outlying Jacobian-norm case. The result was diagnostic rather than a clean mechanistic discovery.",
        "claim": "The sparse-analysis tooling worked, but this branch did not support a strong scientific headline.",
        "source": "reports/jspace_…/analysis_summary.json; commits 6442fff, c07bd6c, 9d54017",
    },
    {
        "title": "5. Pivot from text-only demonstrations to synchronized tri-modal evidence",
        "belief": "A shared workspace hypothesis is most interesting if the same concept is represented across text, image, and spoken language.",
        "trigger": "Text-only evidence could not test cross-modal accessibility.",
        "decision": "Use SpokenCOCO captions synchronized to COCO images, create one photograph/recording per experimental unit, and treat audio as linguistic spoken captions—not environmental sound.",
        "outcome": "The pipeline indexed synchronized groups, audited image leakage, and created disjoint train/test image sets.",
        "claim": "The study tests three input modalities and text output. It does not test image/audio generation or environmental sound.",
        "source": "docs/multimodal_jspace_pilot.md; population manifests and split_provenance.json",
    },
    {
        "title": "6. Select concepts before model results and keep controls fixed",
        "belief": "Concept selection based on outcome strength would inflate evidence.",
        "trigger": "SpokenCOCO contains uneven concept frequency and co-occurrence.",
        "decision": "Rank concepts using dataset coverage, annotation/caption agreement, image uniqueness, and feasibility; freeze six concepts and three focal/control assignments before model stages.",
        "outcome": "The primary population used zebra, cat, toilet, giraffe, bird, and microwave, with unrelated controls assigned deterministically.",
        "claim": "Selection was data-feasibility-driven rather than effect-driven.",
        "source": "mmaudio…/selection_fingerprint and split_provenance; docs/three_modality_claim_admissibility.md",
    },
    {
        "title": "7. Audit native spoken-audio plumbing before interpreting audio",
        "belief": "A valid audio placeholder path and no-op hook audit were prerequisites, but not evidence of concept recognition.",
        "trigger": "The model’s audio interface had different placeholder and feature-count conventions from text/image input.",
        "decision": "Verify placeholder spans, feature agreement, final prompt position, scorer validity for multi-token candidates, activation capture, no-op hooks, and zero interventions.",
        "outcome": "AUDIO_READY passed. The audit explicitly stated that it established technical usability only.",
        "claim": "Engineering evidence only; no audio transfer claim.",
        "source": "audioaudit…/audio_audit.json and audio_audit.md; fingerprint sha256:9ad8…",
    },
    {
        "title": "8. Correct a broken single-token scoring fixture",
        "belief": "Candidate-sequence scoring needed one authoritative path for both one- and multi-token answers.",
        "trigger": "An early fixture made a valid scorer appear invalid and conflated scoring correctness with behavioral capability.",
        "decision": "Return per-token log-probabilities from the same scorer, test multi-token accumulation, and separate scorer validity from whether Gemma answers correctly.",
        "outcome": "The scorer passed structural validity checks; behavioral capability remained a separate empirical gate.",
        "claim": "This repaired measurement plumbing; it did not rescue any behavioral result.",
        "source": "commits d20c741, c30699e, 41317ad; audio audit scorer-validity block",
    },
    {
        "title": "9. Screen behavioral capability before causal interpretation",
        "belief": "A causal effect is uninterpretable if the model cannot identify the concept in the target modality.",
        "trigger": "Spoken-audio accuracy varied by concept even when text and image were strong.",
        "decision": "Require at least 70% accuracy in each required modality. Keep failed concepts in raw tables but prevent them from supporting the principal claim; never replace them post hoc.",
        "outcome": "Five of six concepts passed all three channels. Zebra’s spoken-audio score was 5/8, so zebra was capability-ineligible for the principal tri-modal claim.",
        "claim": "Capability gating narrowed admissible evidence without deleting measured failures.",
        "source": "native_audio_transfer_summary.json; commit dde1bb3; admissibility checksum sha256:8d927…",
    },
    {
        "title": "10. Run the first tri-modal representational and causal study",
        "belief": "A text-calibrated J-space might expose concept structure shared across modalities because all streams enter the same decoder workspace.",
        "trigger": "Capability and engineering gates passed, making the causal test interpretable.",
        "decision": "Measure image-disjoint J-space retrieval and additive residual interventions at L35, then replicate at L38/L40 with matched-random and unrelated controls.",
        "outcome": "Representational retrieval beat shuffled controls for all four audio-related directions. Capability-filtered causal cells supported controlled target effects at L35 and replicated later; the run reported THREE_MODALITY_GO under its predeclared endpoint.",
        "claim": "After the later endpoint audit: evidence for shared candidate-conditioned target log-probability effects across text, image, and spoken captions—not proof that the model freely outputs the swapped concept.",
        "source": "mmaudio…/native_audio_transfer_summary.json; original and capability-filtered reports; endpoint claim ledger",
    },
    {
        "title": "11. Ask whether causal transfer occurred before answer convergence",
        "belief": "The scientifically interesting claim required causal accessibility before the residual already directly named the final answer.",
        "trigger": "L35 was chosen because its lens passed confirmation, not because convergence timing had been measured.",
        "decision": "Audit each stored residual through Gemma’s own final norm and unembedding, with frozen CONVERGED, NOT_CONVERGED, and AMBIGUOUS bars plus permutation controls.",
        "outcome": "L35, L38, and L40 were already converged under this native-readout criterion. The earlier L35 transfer therefore did not support a pre-convergence claim.",
        "claim": "The causal effect existed at or after native direct-readout convergence under this criterion.",
        "source": "commit 49eaad7; docs/output_convergence_timing.md; convergence audit report",
    },
    {
        "title": "12. Extend calibration toward L32 using more text prompts",
        "belief": "L32’s near miss at scale 100 might reflect estimator variance rather than a fundamentally invalid lens.",
        "trigger": "L32 was the closest earlier layer to the validated late-layer region.",
        "decision": "Continue the same running average to 250 and 1,000 nested prompts, use fresh development and untouched confirmation sets, and select the smallest scale matching the largest scale’s eligible set.",
        "outcome": "The development curve plateaued. Scale 250 was selected, and L32 passed that extension’s untouched confirmation set. Later corrected band validation on a different untouched population would overturn L32 eligibility for the band study.",
        "claim": "At this stage only: L32 had one positive independent lens confirmation at scale 250.",
        "source": "rgext…/early_layer_extension_report.json; EARLY_LAYER_CALIBRATION_GO",
    },
    {
        "title": "13. Remove answer choices from causal prompts",
        "belief": "Listing candidate concepts in the prompt could place those concepts into the residual and make intervention effects easier for the wrong reason.",
        "trigger": "Comparison with the paper-style method highlighted that the model should infer the entity from evidence rather than see the candidates named.",
        "decision": "Introduce an open entity-identification prompt, score candidates off-prompt, and add leakage audits preventing audio transcripts or target strings from entering backend arguments.",
        "outcome": "L32 representational transfer remained supported, but causal transfer was weak and not bidirectional. The paired L35 reference under the same open protocol was supported.",
        "claim": "Open-prompt evidence reduced answer-priming concerns but weakened the L32 causal result.",
        "source": "commit 25cdf8a; docs/prompt_protocol.md; mml32…/l32_followup_report.json",
    },
    {
        "title": "14. Resolve L32 convergence on an independent population",
        "belief": "If L32 were demonstrably NOT_CONVERGED, a causal effect there could support the timing claim.",
        "trigger": "The first L32 follow-up classified native readout as AMBIGUOUS.",
        "decision": "Create a separately fingerprinted, photograph-disjoint population with fixed concepts and the same frozen thresholds; run convergence first and gate expensive causal replication on NOT_CONVERGED.",
        "outcome": "The independent population was also AMBIGUOUS: pooled clean-answer agreement 0.611 and target accuracy 0.625. Controls passed. Stage B correctly did not run.",
        "claim": "L32 is neither demonstrated converged nor demonstrated not-converged; ambiguity replicated.",
        "source": "mml32res…/l32_convergence_resolution_summary.json; criterion digest sha256:abbb…",
    },
    {
        "title": "15. Test adjacent L27–L31 with both J-lens and R-lens arms",
        "belief": "An earlier layer might be not-converged while still supporting a faithful coordinate system. R-lens was considered because early-layer J-lens approximation error might be the bottleneck.",
        "trigger": "L26 failed, L32 was ambiguous, and L35 was converged, creating a narrow unresolved interval.",
        "decision": "Fit J- and relevance-weighted R-lens arms at L27–L31 on 250 prompts and open one untouched 256-prompt confirmation set under frozen gates.",
        "outcome": "Neither arm produced a passing layer. The terminal result was BOTH_LENS_ARMS_NO_GO, so no convergence or causal stage was licensed.",
        "claim": "The project did not obtain a confirmed early lens in L27–L31 using either estimator.",
        "source": "mmpre…/adjacent_lens_table.json; commit 867f38a; verdict checksums sha256:106bc… and sha256:a154…",
    },
    {
        "title": "16. Replace additive steering with exact two-coordinate exchange",
        "belief": "The paper’s causal test exchanges source and target coordinates; additive concept-direction steering is a related but different intervention.",
        "trigger": "Method review showed that the earlier causal stages were not an apples-to-apples replication of the paper’s coordinate swap.",
        "decision": "Implement pseudoinverse-based source↔target coordinate exchange, refuse rank-deficient or ill-conditioned bases, recompute coordinates at each patched layer, and distinguish exact α=1 exchange from α=2 sensitivity.",
        "outcome": "The intervention family became faithful to the paper-style algebra. Earlier additive-steering results remained historical evidence and were not relabeled as swaps.",
        "claim": "Method correction, not a retroactive scientific upgrade.",
        "source": "commit 78075ab; docs/coordinate_swap_protocol.md; coordinate_swap.py tests",
    },
    {
        "title": "17. Run a sparse-grid paper-style swap before a full band existed",
        "belief": "Existing validated layers L32/L35/L38/L40 could provide an initial timing diagnostic.",
        "trigger": "The exact exchange implementation was ready, but interior layer lenses were not.",
        "decision": "Run both hidden-intermediate and direct-answer arms across the sampled grid while explicitly labeling it a grid, not a contiguous band.",
        "outcome": "The first run was inconclusive. A revised v2 produced an α=1 intermediate-null/answer-positive pattern and an α=2 earlier-intermediate sensitivity in one direction, but independent α=2 confirmation did not reproduce the intermediate effect.",
        "claim": "Exploratory timing evidence only. The α=2 intermediate result did not independently replicate.",
        "source": "mmpaper…, mmpaper2…, and both mmpaperconfirm… reports; direction-matched amendment v3",
    },
    {
        "title": "18. Enforce the paper’s contiguous-band requirement",
        "belief": "A band intervention must patch every physical layer in the interval; a sparse grid cannot stand in for a band.",
        "trigger": "The tested grid omitted L33, L34, L36, L37, and L39.",
        "decision": "Fit the five missing scale-250 lenses and block Stage 3 unless every interior layer passed untouched confirmation.",
        "outcome": "The initial interior confirmation reported no passing layers, which led to an audit of the wrong-layer control.",
        "claim": "No band result was allowed while lens validity was unresolved.",
        "source": "commit c28e027; bandlens…/band_interior_lens_report.json",
    },
    {
        "title": "19. Repair a set-dependent wrong-layer control",
        "belief": "A wrong-layer control should be meaningfully distant and fixed independently of whichever subset was newly fitted.",
        "trigger": "The initial control compared L33/L34 with L39 and L36–L39 with L33, so the ‘wrong’ lens was another strong nearby late-layer lens.",
        "decision": "Freeze a control universe including L8/L14/L20/L26 and use L8 as the distant source for L32–L40. Re-run confirmation without refitting or changing thresholds.",
        "outcome": "L33–L40 passed; L32 failed coverage/nondegeneracy. The originally planned L32–L40 band remained permanently NO_GO, and the largest admissible band became L33–L40.",
        "claim": "The correction rescued the interior lenses but not L32 or the original confirmatory design.",
        "source": "commit cd89bc9; bandcorr…/band_interior_corrected_validation_report.json",
    },
    {
        "title": "20. Run a prospective L33–L40 band follow-up",
        "belief": "A fully validated contiguous late-layer band could still test whether coordinate exchange changes hidden and answer endpoints.",
        "trigger": "Corrected confirmation established L33–L40 only after lens results were seen.",
        "decision": "Run the band as a prospective follow-up, keep the original L32–L40 NO_GO unchanged, and report α=1 primary separately from α=2 sensitivity.",
        "outcome": "No α=1 primary band passed. The full L33–L40 band appeared only in the α=2 sensitivity arm. Intermediate effects were null; answer effects appeared at deeper starts; timing remained inconclusive.",
        "claim": "α=2 sensitivity evidence in a prospectively selected band, not confirmation of the original band or a timing onset.",
        "source": "commit cc4df53; band3340…/l33_l40_validated_band_followup_report.json",
    },
    {
        "title": "21. Audit endpoint semantics across every active claim",
        "belief": "A logit increase, a restricted-candidate win, and an unrestricted generated token are not interchangeable endpoints.",
        "trigger": "Teacher-forced and candidate-restricted scoring had been described too loosely as changing the model’s answer.",
        "decision": "Trace 24 endpoints through code, label each as representational, conditional sequence log-probability, restricted candidate, unrestricted vocabulary, greedy demonstration, or engineering, and narrow claim language where needed.",
        "outcome": "Thirteen claims survived unchanged; eleven survived with narrower wording; zero were discarded as wholly unsupported. Crucially, the early tri-modal causal result became a controlled target-log-probability claim.",
        "claim": "The endpoint audit changed interpretation, not the stored measurements.",
        "source": "commit 2cba864; reports/endpoint_audit/endpoint_semantics_audit.md and endpoint_claim_ledger.json",
    },
    {
        "title": "22. Re-run the band study with unrestricted vocabulary scoring",
        "belief": "A stronger behavioral claim required the target to win against the entire vocabulary rather than a supplied candidate set.",
        "trigger": "The endpoint audit showed that restricted preference could not establish free model output.",
        "decision": "Reuse the same population and completed band evidence, score unrestricted next-token output, add tokenization requirements, and preserve completed artifacts as read-only.",
        "outcome": "The unrestricted word-token study returned FULL_VOCAB_REASONING_NO_GO; the cross-modal conjunction was inconclusive.",
        "claim": "The restricted α=2 effect did not translate into the stronger full-vocabulary endpoint.",
        "source": "mmfv…/full_vocabulary_causal_validation_report.json; report checksum sha256:669a…",
    },
    {
        "title": "23. Move from animal words to single-token digit reasoning",
        "belief": "The downstream leg-count answer offered a clean categorical endpoint: bird→cat should change 2→4 and cat→bird should change 4→2.",
        "trigger": "Word candidates and prompt formatting complicated unrestricted endpoint interpretation.",
        "decision": "Audit tokenizer behavior, fix the initial false multi-token refusal, freeze single-token digit endpoints, and run hidden-intermediate and direct-answer arms with full-vocabulary greedy parity.",
        "outcome": "The α=2 digit confirmation was NO_GO even though clean digit capability passed.",
        "claim": "No unrestricted downstream reasoning confirmation at α=2.",
        "source": "commits 0ddc4c8 and 92e393e; mmdigitconfirm…/digit_reasoning_confirmation_report.json",
    },
    {
        "title": "24. Return to exact α=1 exchange and expand recruitment",
        "belief": "Alpha should not be tuned to force a preferred result; exact exchange at α=1 is the paper-comparable primary intervention. More source images could improve the chance of recruiting clean, capable cells without changing the endpoint.",
        "trigger": "The first exact α=1 run failed its capability screen before interventions.",
        "decision": "Keep α=1 fixed, expand the recruitment pool to 64 images per concept, select eight clean examples per direction × modality cell, and retain zero, random, unrelated, and direct-answer controls.",
        "outcome": "Hidden-intermediate success was 2/48 and did not beat paired controls. The bird→cat direct-answer control was 24/24 across modalities; cat→bird was 0/24. Pooled direct-answer effects beat zero and random controls, proving endpoint reachability in one direction but exposing strong asymmetry.",
        "claim": "The strongest unrestricted result is a directional positive control, not successful downstream reasoning transfer.",
        "source": "commits c8a7d53 and b643e1a; mmalpha1confirm64… report checksum sha256:29b969…",
    },
    {
        "title": "25. Freeze the final interpretation instead of tuning toward a positive headline",
        "belief": "A rigorous application is stronger when it distinguishes useful positive evidence from a failed headline hypothesis.",
        "trigger": "The exact paper-style, unrestricted endpoint did not support the desired hidden-reasoning claim.",
        "decision": "Stop increasing α or changing thresholds after outcomes; preserve the positive candidate-conditioned tri-modal result, the validated lens map, the endpoint audit, and the one-direction direct-answer control as separate findings.",
        "outcome": "The project ends with a clear mechanistic diagnosis: the intervention can strongly move an exposed answer coordinate in bird→cat, but the text-calibrated J-lens coordinate exchange does not reliably recruit downstream hidden reasoning across modalities.",
        "claim": "A bounded, honest set of findings plus a concrete next-design hypothesis—not the originally desired strong claim.",
        "source": "final archived reports and endpoint ledger; no post-result threshold changes",
    },
]


RUNS = [
    ("rgcalib_real_7e3736b4de8f", "Text J-lens calibration", "L35/L38/L40 pass at scale 100"),
    ("rgext_real_c18f03f06e7b", "Early-layer extension", "L32 passes one scale-250 confirmation; later superseded for band eligibility"),
    ("audioaudit_real_20260805T012451_c2d58028", "Native-audio engineering audit", "AUDIO_READY; no concept-transfer claim"),
    ("mmaudio_native_audio_transfer_20260806T144822", "Tri-modal candidate-conditioned study", "Representational and controlled causal effects supported"),
    ("mml32_l32_followup_20260808T182717", "Open-prompt L32 follow-up", "Causal WEAK; convergence AMBIGUOUS"),
    ("mml32res_l32_convergence_resolution_20260810T174731", "Independent L32 resolution", "AMBIGUOUS replicated"),
    ("mmpre_real_fdddd750b4c0", "L27–L31 J-/R-lens study", "BOTH_LENS_ARMS_NO_GO"),
    ("mmpaper_real_24be1d028bf1", "Paper-style swap v1", "INCONCLUSIVE"),
    ("mmpaper2_real_04ab55235502", "Paper-style swap v2", "α=2 sensitivity appeared in one direction"),
    ("mmpaperconfirm_real_6b0745c08d84", "First α=2 confirmation", "CAPABILITY_NO_GO"),
    ("mmpaperconfirm_real_a496d5ad7f18", "Independent α=2 confirmation", "Intermediate effect did not replicate"),
    ("bandlens_real_de9338ec2a6e", "Interior lens validation", "Initial NO_GO under defective nearby wrong-layer control"),
    ("bandcorr_real_eb5b00f135e4", "Corrected lens validation", "L33–L40 pass; L32 fails"),
    ("band3340_real_2a72bda9b4ba", "Prospective L33–L40 band", "α=2 sensitivity only; timing inconclusive"),
    ("mmfv_real_bfb07903e961", "Full-vocabulary word endpoint", "NO_GO"),
    ("mmdigitconfirm_real_68c182bfc025", "Digit endpoint α=2", "NO_GO"),
    ("mmalpha1confirm_real_6212d73dc72a", "Exact α=1 first attempt", "CAPABILITY_NO_GO"),
    ("mmalpha1confirm64_real_df0ce0404c32", "Exact α=1 recruitment-64", "Hidden reasoning NO_GO; directional answer control positive"),
]


def build() -> Path:
    doc = Document()
    configure(doc)
    add_title(doc, "Multimodal J-Space Research Decision Log")
    add_subtitle(doc, "Rationale, evidence, pivots, and claim boundaries · July 14–August 18, 2026")
    add_callout(
        doc,
        "Evidence standard",
        "This document was reconstructed from repository history and 18 archived run directories. Console recollections are not treated as results. Superseded and negative runs remain visible because they explain why the design changed.",
        "blue",
    )
    p = doc.add_paragraph()
    p.add_run("Purpose. ").bold = True
    p.add_run(
        "Provide a first-principles account of what the project tried, why each decision was made, what was actually measured, what changed after audits, and what can be defended in a MATS application."
    )

    doc.add_heading("Contents", level=1)
    for item in [
        "1. Executive summary",
        "2. First-principles method",
        "3. Claim ladder and endpoint semantics",
        "4. Chronological decision and pivot ledger",
        "5. What the evidence now says",
        "6. Methodological audit",
        "7. A cleaner next research design",
        "Appendix A. Archived run manifest",
        "Appendix B. Reproducible figure notebook",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    page_break(doc)
    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "The project began with a broad question: does Gemma 4 use a modality-general workspace whose verbalizable concept coordinates can be read and causally exchanged across text, images, and spoken captions? The answer is mixed. The work produced real positive evidence, but not the strongest paper-style claim."
    )
    add_callout(
        doc,
        "Strongest positive finding",
        "A text-calibrated J-space showed high image-disjoint retrieval across text, image, and spoken captions, and controlled interventions increased target-candidate log-probability in multiple cross-modal directions. This supports shared candidate-conditioned concept accessibility—not unrestricted generation of the swapped concept.",
        "green",
    )
    add_callout(
        doc,
        "Strongest unrestricted causal observation",
        "Under exact α=1 coordinate exchange and unrestricted greedy digit output, the direct-answer bird→cat control changed 2→4 on 24/24 selected examples across all three modalities. The reverse direction was 0/24, and the hidden-intermediate arm was only 2/48. The mechanism can move one exposed answer endpoint, but it did not reliably alter downstream reasoning.",
        "amber",
    )
    add_callout(
        doc,
        "Headline hypothesis status",
        "Not confirmed. The project did not demonstrate reliable, bidirectional, cross-modal downstream reasoning transfer before answer convergence under the exact paper-style unrestricted endpoint.",
        "red",
    )
    doc.add_paragraph(
        "The research value lies in the sequence of falsifiable repairs: independent confirmation sets, capability admissibility, prompt decontamination, native-head convergence audits, exact coordinate exchange, contiguous-band enforcement, wrong-layer-control repair, endpoint-semantics auditing, and unrestricted-vocabulary confirmation. Each repair narrowed the claim rather than hiding an unfavorable outcome."
    )
    add_picture(doc, "figure_3_endpoint_pivot_timeline.png", "Figure 1. Tightening the endpoint repeatedly reduced the strength of the scientific claim while increasing its behavioral meaning.")

    page_break(doc)
    doc.add_heading("2. First-principles method", level=1)
    doc.add_heading("2.1 The object being estimated", level=2)
    doc.add_paragraph(
        "At a chosen transformer layer ℓ, let hℓ be the residual-stream activation at the final prompt token. Gemma’s remaining computation maps hℓ to output logits. The J-lens estimates the average Jacobian of that mapping over a text calibration population. Its columns define approximate verbalizable directions: directions that, under the local linear approximation, affect output-token logits."
    )
    add_callout(
        doc,
        "Important limitation",
        "The lens is a population-average linear approximation, not the model’s true ontology. It can miss concepts without clean token names, blur nonlinear structure, and amplify estimation error during causal intervention.",
        "amber",
    )
    doc.add_heading("2.2 Why a text-calibrated lens was used across modalities", level=2)
    doc.add_paragraph(
        "The original scientific bet was intentionally stronger than fitting one lens per modality: if image and spoken-caption activations can be decoded through the same text-calibrated coordinate system, that is evidence of a shared verbalizable workspace. A tri-modal-fitted lens would likely improve reconstruction but would answer a different question and risk learning the alignment being tested. The text-only fit is therefore defensible for representational testing, but it may be a weak instrument for causal exchange."
    )
    doc.add_heading("2.3 Three distinct intervention families", level=2)
    add_table(
        doc,
        ["Family", "Operation", "What it can establish"],
        [
            ["Additive direction steering", "Add/subtract a learned J-space concept direction at one layer", "Controlled sensitivity; not an exact paper-style swap"],
            ["Two-coordinate exchange", "Solve source/target coordinates with a pseudoinverse and exchange them at α=1", "Paper-comparable coordinate manipulation at a layer"],
            ["Contiguous band clamp", "Repeat the exchange at every physical layer in a validated interval", "Persistence of a coordinate through a layer band; starts are not exact onsets"],
        ],
        widths=[1.65, 2.55, 2.3],
    )
    page_break(doc)
    doc.add_heading("2.4 Four endpoint classes", level=2)
    add_table(
        doc,
        ["Endpoint", "Question answered", "Strength"],
        [
            ["Representational retrieval", "Are matching concepts close in J-space across modalities?", "Correlational structure"],
            ["Conditional sequence log-probability", "Does the intervention increase probability of a named candidate sequence?", "Controlled causal preference"],
            ["Restricted-candidate winner", "Does the target beat a small supplied candidate set?", "Behavioral but candidate-conditioned"],
            ["Unrestricted greedy token", "Does the target win against the whole vocabulary without an appended answer?", "Strongest endpoint used here"],
        ],
        widths=[1.6, 3.1, 1.8],
    )

    doc.add_heading("3. Claim ladder and endpoint semantics", level=1)
    doc.add_paragraph(
        "The project repeatedly appeared to approach a strong result because several weaker endpoints were positive. The endpoint audit made the hierarchy explicit. Passing a lower rung does not imply passing a higher one."
    )
    for text in [
        "The model can identify the source concept in each modality.",
        "Cross-modal activations are close under a frozen text-calibrated J-space.",
        "An intervention changes a target candidate’s conditional score against controls.",
        "The target wins a restricted candidate set.",
        "The target becomes the unrestricted next token.",
        "Changing the hidden concept causes a downstream answer to change reliably and bidirectionally.",
    ]:
        doc.add_paragraph(text, style="List Number")
    add_callout(
        doc,
        "Final position on the ladder",
        "The project reached rung 3 robustly for the original tri-modal study, reached rung 5 strongly only for a one-direction direct-answer positive control, and did not reach rung 6 for hidden downstream reasoning.",
        "blue",
    )

    page_break(doc)
    doc.add_heading("4. Chronological decision and pivot ledger", level=1)
    doc.add_paragraph(
        "Each entry separates what was believed at the time from hindsight. ‘Outcome’ reports what the archived artifacts say; ‘claim impact’ states how interpretation changed."
    )
    for pivot in PIVOTS:
        doc.add_heading(pivot["title"], level=2)
        for label, key in [
            ("Starting belief", "belief"),
            ("Trigger", "trigger"),
            ("Decision", "decision"),
            ("Measured outcome", "outcome"),
            ("Claim impact", "claim"),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(label + ". ")
            r.bold = True
            r.font.name = "Arial"
            r.font.color.rgb = RGBColor.from_string(BLACK)
            p.add_run(pivot[key])
        add_source(doc, pivot["source"])

    doc.add_heading("5. What the evidence now says", level=1)
    doc.add_heading("5.1 Lens validity is strongly depth-dependent", level=2)
    doc.add_paragraph(
        "The corrected common-population confirmation shows a clean late-layer trend. L33–L40 passed; L32 failed. Earlier L27–L31 failed in both J- and R-lens arms. This makes the lens most trustworthy after much of the answer structure is already available, which is exactly where a pre-convergence claim becomes difficult."
    )
    add_picture(doc, "figure_1_corrected_lens_validation.png", "Figure 2. Corrected untouched confirmation mean reciprocal rank. Color denotes the frozen pass/fail decision, not an MRR threshold.")
    doc.add_heading("5.2 Behavioral capability was mostly strong", level=2)
    doc.add_paragraph(
        "The original selected concepts were easy for Gemma in text and image. Spoken captions were also strong except zebra, which failed the 70% gate. This matters because the later null result cannot be dismissed as a general inability to process speech, but concept-level capability still constrained admissible cells."
    )
    add_picture(doc, "figure_2_trimodal_capability.png", "Figure 3. Restricted six-candidate behavioral capability. Zebra spoken audio is the only cell below the predeclared 70% gate.", width=5.35)
    doc.add_heading("5.3 L32 ambiguity replicated", level=2)
    doc.add_paragraph(
        "The independent L32 population landed between the frozen NOT_CONVERGED and CONVERGED bars. This is not evidence that L32 converged. It is evidence that the chosen linear native-readout test could not classify it cleanly. More importantly, the project never found a confirmed earlier lens that would let it move the same causal test to a demonstrably not-converged layer."
    )
    add_picture(doc, "figure_5_l32_ambiguity.png", "Figure 4. Pooled independent-population metrics. The complete criterion required every modality and additional rank conditions; the displayed values illustrate why neither bar was crossed.")
    doc.add_heading("5.4 Exact exchange reveals instrument asymmetry", level=2)
    doc.add_paragraph(
        "The final α=1 study is the cleanest behavioral audit. Its hidden-intermediate arm rarely changed the downstream digit. The direct-answer arm, however, perfectly changed bird→cat across text, image, and spoken captions while never changing cat→bird. That asymmetry points to geometry or endpoint accessibility—not a universally ineffective hook."
    )
    add_picture(doc, "figure_4_alpha1_unrestricted_outcomes.png", "Figure 5. Unrestricted greedy next-token successes out of eight in each direction × modality cell.")

    doc.add_heading("5.5 Defensible findings", level=2)
    for finding in [
        "A research-grade Gemma 4 J-lens pipeline was implemented with frozen architecture checks, independent confirmation, provenance fingerprints, and resumable units.",
        "A frozen text-calibrated J-space supports high image-disjoint retrieval among text, images, and spoken captions at validated late layers.",
        "Controlled cross-modal interventions change target candidate probabilities against matched controls; this is a candidate-conditioned causal result.",
        "L32 native readout was ambiguous on two populations, while L33–L40 lenses passed corrected confirmation and L27–L31 lenses did not.",
        "Exact α=1 coordinate exchange can strongly move a direct answer in one direction across all three input modalities, but did not reliably produce the downstream hidden-reasoning effect.",
        "Endpoint semantics materially changed interpretation: restricted and teacher-forced evidence must not be reported as free generation.",
    ]:
        doc.add_paragraph(finding, style="List Bullet")

    doc.add_heading("5.6 Claims this evidence does not support", level=2)
    for claim in [
        "Gemma freely outputs the injected concept after the original tri-modal intervention.",
        "The same hidden concept is causally used for downstream reasoning before answer convergence.",
        "The effect is bidirectional, symmetric, or robust to all source–target pairs.",
        "The model shares one modality-independent representation at every layer.",
        "Environmental audio participates in the result.",
        "A tri-modal-fitted lens would necessarily solve the causal problem.",
    ]:
        doc.add_paragraph(claim, style="List Bullet")

    doc.add_heading("6. Methodological audit", level=1)
    doc.add_heading("6.1 Flaws found and repaired", level=2)
    add_table(
        doc,
        ["Issue", "Repair", "Effect on interpretation"],
        [
            ["Candidate answers present in prompts", "Open prompts + leakage refusal", "Reduced target priming; L32 causal result weakened"],
            ["Additive steering described near swap language", "Exact pseudoinverse coordinate exchange", "Separated historical steering from paper-style causal evidence"],
            ["Sparse grid called a band", "Fit and validate all interior layers", "Blocked claims until a true contiguous interval existed"],
            ["Nearby late-layer wrong control", "Fixed universe with distant L8 control", "Validated L33–L40; rejected L32"],
            ["Teacher-forced/restricted endpoints overstated", "24-claim endpoint audit + full-vocabulary reruns", "Narrowed eleven claims and produced NO_GO at stronger endpoints"],
            ["Digit tokenizer falsely treated as multi-token", "Tokenizer-aware endpoint resolver", "Enabled the unrestricted digit study"],
            ["Post-result concept substitution risk", "Capability-ineligible concepts retained but excluded", "Prevented outcome-driven replacement"],
        ],
        widths=[1.75, 2.2, 2.55],
    )
    doc.add_heading("6.2 Remaining methodological risks", level=2)
    add_table(
        doc,
        ["Risk", "Why it matters", "Best next test"],
        [
            ["Text-calibrated average Jacobian is noisy", "Causal exchange may inject error even when readout works", "Compare J-lens, R-lens, supervised probe, and SAE directions on the same frozen task"],
            ["Only one animal pair is deeply confirmed", "Bird→cat / cat→bird asymmetry may be pair-specific", "Predeclare several single-token, capability-matched pairs"],
            ["Prompt endpoint may recruit shallow heuristics", "Leg-count output can be solved without a stable hidden animal state", "Use multi-step tasks with hidden intermediate labels and held-out downstream properties"],
            ["Shared decoder does not imply shared coordinates", "Modality encoders may deliver aligned but differently scaled states", "Measure modality-conditioned coordinate calibration without fitting on outcome data"],
            ["Band starts are nested", "An effective suffix does not identify exact physical onset", "Use single-layer clamps or non-nested windows with predeclared comparisons"],
            ["Selection and repairs accumulated over time", "Later studies are prospective follow-ups, not one untouched confirmatory sequence", "Freeze one compact protocol and run it once on a fresh population"],
        ],
        widths=[1.65, 2.35, 2.5],
    )
    doc.add_heading("6.3 Hindsight: what consumed the most time", level=2)
    doc.add_paragraph(
        "The project optimized implementation before freezing the strongest behavioral endpoint. This led to repeated expensive reruns when prompt semantics, swap fidelity, band geometry, control geometry, and tokenization were clarified. The largest process improvement is therefore not another model run; it is to define the final behavioral endpoint, positive control, null controls, admissibility rule, and claim sentence before fitting or intervening."
    )

    doc.add_heading("7. A cleaner next research design", level=1)
    doc.add_paragraph(
        "The next project should treat the current work as a measurement audit and instrument-diagnosis study. It should not try to rescue the existing headline by tuning α. A fresh design can test why the direct-answer coordinate is reachable but the hidden reasoning coordinate is not."
    )
    for title, text in [
        ("Freeze the claim", "‘Exchanging a hidden entity coordinate changes at least two independently scored downstream properties under unrestricted output, beyond matched zero/random/unrelated controls.’"),
        ("Choose the model by task capability", "Screen several manageable multimodal models and select before intervention using only clean-task capability, tokenizer simplicity, and accessible hooks—not causal effect size."),
        ("Choose concepts by endpoint geometry", "Use several pairs with single-token entity and property answers, balanced clean margins, and both directions behaviorally capable."),
        ("Compare instruments", "Fit a text J-lens, modality-conditioned J-lens, R-lens, supervised linear probe, and possibly SAE direction on identical frozen activations. Evaluate each on held-out readout and intervention specificity."),
        ("Separate observation from intervention", "First test whether each direction reads the hidden entity. Then test exact exchange at α=1. Keep α sweeps explicitly secondary and multiplicity-corrected."),
        ("Use a fresh confirmatory population", "After selecting one instrument and one task on development data, lock code, thresholds, pairs, layers, and sample count; run once on untouched images/recordings."),
    ]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)
    add_callout(
        doc,
        "Recommended application framing",
        "Present this project as evidence of research judgment under pressure: you built a difficult multimodal causal pipeline, discovered that your original endpoint was weaker than your language, audited it, repaired the method toward the public paper, and learned exactly where the strong claim failed. The core contribution is not a forced positive headline; it is a transparent map from representational evidence to behavioral causal failure.",
        "green",
    )

    page_break(doc)
    doc.add_heading("Appendix A. Archived run manifest", level=1)
    doc.add_paragraph(
        "The following 18 Drive archives were provided and inspected. Binary lens tensors were not needed for the reporting analysis; report JSON, Markdown, CSV/JSONL, fingerprints, and unit summaries were extracted."
    )
    add_table(doc, ["Run", "Purpose", "Recorded outcome"], [list(row) for row in RUNS], widths=[2.35, 1.9, 2.25])

    page_break(doc)
    add_explicit_h1(doc, "Appendix B. Reproducible figure notebook")
    doc.add_paragraph(
        "Notebook: notebooks/mats_research_evidence_analysis.ipynb\nAnalysis module: scripts/mats_research_evidence.py\nMachine-readable output: reports/mats_application/figures/evidence_summary.json"
    )
    doc.add_paragraph(
        "The notebook reads archived reports and regenerates Figures 1–5. It does not load Gemma, fit a lens, execute an intervention, or change a verdict. Missing evidence produces a hard failure."
    )
    add_explicit_h1(doc, "Appendix C. Core repository sources")
    for source in [
        "docs/research_log.md",
        "docs/paper_outline.md",
        "docs/multimodal_jspace_pilot.md",
        "docs/coordinate_swap_protocol.md",
        "docs/endpoint_semantics.md",
        "docs/three_modality_claim_admissibility.md",
        "docs/prompt_protocol.md",
        "docs/research_grade_early_layer_extension_protocol.md",
        "docs/output_convergence_timing.md",
        "docs/l27_l31_preconvergence_study.md",
        "reports/endpoint_audit/endpoint_semantics_audit.md",
        "reports/endpoint_audit/endpoint_claim_ledger.json",
    ]:
        doc.add_paragraph(source, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
